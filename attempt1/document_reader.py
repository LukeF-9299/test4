#!/usr/bin/env python3
"""
Document reader script for reading various file types including .doc, .docx, PDF, and .txt files.
"""

import os
import subprocess
from pathlib import Path
from typing import Optional, Union

try:
    from docx import Document
except ImportError:
    print("Warning: python-docx not installed. Run: pip install python-docx")
    Document = None

try:
    import fitz  # PyMuPDF
except ImportError:
    print("Warning: PyMuPDF not installed. Run: pip install pymupdf")
    fitz = None


def read_doc_file(file_path: Union[str, Path]) -> Optional[str]:
    """
    Read .doc file using antiword command-line tool.
    
    Args:
        file_path: Path to the .doc file
        
    Returns:
        Text content of the file or None if failed
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        print(f"File not found: {file_path}")
        return None
    
    if file_path.suffix.lower() != '.doc':
        print(f"Expected .doc file, got: {file_path.suffix}")
        return None
    
    try:
        # Use antiword to extract text from .doc file
        result = subprocess.run(
            ['antiword', str(file_path)],
            capture_output=True,
            text=True,
            check=True
        )
        return result.stdout
    except subprocess.CalledProcessError as e:
        print(f"Error reading .doc file with antiword: {e}")
        print(f"Stderr: {e.stderr}")
        return None
    except FileNotFoundError:
        print("Error: antiword not found. Please install antiword:")
        print("  Ubuntu/Debian: sudo apt-get install antiword")
        print("  macOS: brew install antiword")
        return None


def read_docx_file(file_path: Union[str, Path]) -> Optional[str]:
    """
    Read .docx file using python-docx library.
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        Text content of the file or None if failed
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        print(f"File not found: {file_path}")
        return None
    
    if file_path.suffix.lower() != '.docx':
        print(f"Expected .docx file, got: {file_path.suffix}")
        return None
    
    if Document is None:
        print("Error: python-docx not installed")
        return None
    
    try:
        doc = Document(str(file_path))
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return '\n'.join(text_content)
    except Exception as e:
        print(f"Error reading .docx file: {e}")
        return None


def read_pdf_file(file_path: Union[str, Path]) -> Optional[str]:
    """
    Read PDF file using PyMuPDF library.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Text content of the file or None if failed
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        print(f"File not found: {file_path}")
        return None
    
    if file_path.suffix.lower() != '.pdf':
        print(f"Expected .pdf file, got: {file_path.suffix}")
        return None
    
    if fitz is None:
        print("Error: PyMuPDF not installed")
        return None
    
    try:
        doc = fitz.open(str(file_path))
        text_content = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            if text.strip():
                text_content.append(f"--- Page {page_num + 1} ---")
                text_content.append(text.strip())
        
        doc.close()
        return '\n'.join(text_content)
    except Exception as e:
        print(f"Error reading PDF file: {e}")
        return None


def read_txt_file(file_path: Union[str, Path]) -> Optional[str]:
    """
    Read .txt file directly.
    
    Args:
        file_path: Path to the .txt file
        
    Returns:
        Text content of the file or None if failed
    """
    file_path = Path(file_path)
    
    if not file_path.exists():
        print(f"File not found: {file_path}")
        return None
    
    if file_path.suffix.lower() != '.txt':
        print(f"Expected .txt file, got: {file_path.suffix}")
        return None
    
    try:
        # Read text file with UTF-8 encoding, fallback to latin-1 if needed
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    except Exception as e:
        print(f"Error reading .txt file: {e}")
        return None


def read_document(file_path: Union[str, Path]) -> Optional[str]:
    """
    Read document file based on its extension.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Text content of the file or None if failed
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()
    
    if suffix == '.doc':
        return read_doc_file(file_path)
    elif suffix == '.docx':
        return read_docx_file(file_path)
    elif suffix == '.pdf':
        return read_pdf_file(file_path)
    elif suffix == '.txt':
        return read_txt_file(file_path)
    else:
        print(f"Unsupported file type: {suffix}")
        return None


def read_storage_folder(storage_path: Union[str, Path] = "storage") -> dict:
    """
    Read all supported documents in the storage folder.
    
    Args:
        storage_path: Path to the storage folder
        
    Returns:
        Dictionary with file paths as keys and text content as values
    """
    storage_path = Path(storage_path)
    
    if not storage_path.exists():
        print(f"Storage folder not found: {storage_path}")
        return {}
    
    if not storage_path.is_dir():
        print(f"Storage path is not a directory: {storage_path}")
        return {}
    
    supported_extensions = {'.doc', '.docx', '.pdf', '.txt'}
    results = {}
    
    for file_path in storage_path.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            print(f"Reading: {file_path}")
            content = read_document(file_path)
            if content:
                results[str(file_path)] = content
    
    return results


if __name__ == "__main__":
    # Example usage
    storage_folder = "storage"
    documents = read_storage_folder(storage_folder)
    
    print(f"\nFound {len(documents)} documents:")
    for file_path, content in documents.items():
        print(f"- {file_path}: {len(content)} characters")
